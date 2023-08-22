# Challenge #2
# https://github.com/Robpol86/terminaltables
# 1. Use pipenv to install terminaltables
# 2. Read the online documentation to get the example terminal table working
from terminaltables import AsciiTable
table_data = [
    ['Heading1', 'Heading2'],
    ['row1 column1', 'row1 column2'],
    ['row2 column1', 'row2 column2'],
    ['row3 column1', 'row3 column2']
]
table = AsciiTable(table_data)
print(table.table)

# Challenge #3
# https://pypi.org/project/python-docx/
# 1. Use pipenv to install pyton-docx
# 2. Read the online documentation to get reading and writing to the
# example.docx working
from docx import Document
from docx.shared import Inches


# found this:
# https://python-docx.readthedocs.io/en/latest/user/documents.html
document = Document('example.docx')

# And used code from here:
# https://python-docx.readthedocs.io/en/latest/
# for the rest
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('new-file-name.docx')

